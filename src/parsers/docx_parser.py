"""
DOCX Parser - Extract text from Word documents.
"""
from typing import Optional, List
from pathlib import Path
import re


class DocxParser:
    """Parse DOCX files and extract text content."""
    
    def __init__(self):
        """Initialize DOCX parser."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are installed."""
        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    def parse(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"Not a Word file: {file_path}")
        
        try:
            doc = self.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            text = '\n'.join(paragraphs)
            return self._clean_text(text)
            
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {e}")
    
    def parse_bytes(self, docx_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes.
        
        Args:
            docx_bytes: DOCX file content as bytes
            
        Returns:
            Extracted text content
        """
        import io
        
        try:
            doc = self.Document(io.BytesIO(docx_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            text = '\n'.join(paragraphs)
            return self._clean_text(text)
            
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX bytes: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Clean up
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
